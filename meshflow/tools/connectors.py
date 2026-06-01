"""MeshFlow curated connector library — 50 pre-built @tool-decorated async functions.

Covers 10 SaaS categories with ~5 tools each:
  Communication, GitHub, Web Search/Scraping, Data/Storage, Productivity,
  AI/LLM Utilities, CRM/Sales, Calendar/Meetings, Finance/Data, DevOps/Infrastructure,
  File/Document Processing.

All HTTP calls use urllib.request (zero extra dependencies).
Credentials are read from environment variables.
Errors are returned as strings rather than raised.

Usage::

    from meshflow.tools.connectors import TOOL_REGISTRY
    tool = TOOL_REGISTRY["slack_post_message"]
    result = await tool.call(channel="#general", text="Hello from MeshFlow!")
"""

from __future__ import annotations

import base64
import json
import os
import urllib.error
import urllib.parse
import urllib.request
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from smtplib import SMTP, SMTP_SSL
from typing import Any

from meshflow.core.schemas import RiskTier
from meshflow.tools.registry import tool


# ── Shared HTTP helpers ────────────────────────────────────────────────────────


def _http_get(url: str, headers: dict[str, str] | None = None) -> Any:
    """Make a GET request and return parsed JSON or raw text."""
    req = urllib.request.Request(url, headers=headers or {})
    try:
        with urllib.request.urlopen(req, timeout=15) as resp:
            raw = resp.read()
            ct = resp.headers.get("Content-Type", "")
            if "json" in ct or raw.lstrip().startswith(b"{") or raw.lstrip().startswith(b"["):
                return json.loads(raw)
            return raw.decode("utf-8", errors="replace")
    except urllib.error.HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")
        return {"error": f"HTTP {exc.code}: {body[:400]}"}
    except Exception as exc:
        return {"error": str(exc)}


def _http_post(
    url: str,
    payload: Any,
    headers: dict[str, str] | None = None,
    method: str = "POST",
) -> Any:
    """Make a POST/PATCH/PUT request with JSON body and return parsed JSON."""
    data = json.dumps(payload).encode()
    hdrs = {"Content-Type": "application/json", **(headers or {})}
    req = urllib.request.Request(url, data=data, headers=hdrs, method=method)
    try:
        with urllib.request.urlopen(req, timeout=15) as resp:
            raw = resp.read()
            if raw:
                return json.loads(raw)
            return {"ok": True}
    except urllib.error.HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")
        return {"error": f"HTTP {exc.code}: {body[:400]}"}
    except Exception as exc:
        return {"error": str(exc)}


def _bearer(token: str) -> dict[str, str]:
    return {"Authorization": f"Bearer {token}"}


# ══════════════════════════════════════════════════════════════════════════════
# CATEGORY 1 — COMMUNICATION
# ══════════════════════════════════════════════════════════════════════════════


@tool(
    name="slack_post_message",
    description="Post a message to a Slack channel. Requires SLACK_BOT_TOKEN.",
    risk=RiskTier.EXTERNAL_IO,
    tags=["communication", "slack"],
)
async def slack_post_message(channel: str, text: str) -> str:
    """Send a text message to a Slack channel or DM.

    channel: Slack channel ID or name (e.g. '#general' or 'C01234ABC').
    text: Message body (plain text or mrkdwn).
    """
    token = os.environ.get("SLACK_BOT_TOKEN", "")
    if not token:
        return "Error: SLACK_BOT_TOKEN not set"
    result = _http_post(
        "https://slack.com/api/chat.postMessage",
        {"channel": channel, "text": text},
        headers=_bearer(token),
    )
    if isinstance(result, dict) and not result.get("ok"):
        return f"Slack error: {result.get('error', result)}"
    return f"Message sent to {channel} (ts={result.get('ts', 'unknown')})"


@tool(
    name="slack_get_channel_history",
    description="Retrieve recent messages from a Slack channel. Requires SLACK_BOT_TOKEN.",
    risk=RiskTier.READ_ONLY,
    tags=["communication", "slack"],
)
async def slack_get_channel_history(channel: str, limit: int = 10) -> str:
    """Fetch the most recent messages from a Slack channel.

    channel: Channel ID (e.g. 'C01234ABC').
    limit: Number of messages to return (default 10, max 1000).
    """
    token = os.environ.get("SLACK_BOT_TOKEN", "")
    if not token:
        return "Error: SLACK_BOT_TOKEN not set"
    url = (
        f"https://slack.com/api/conversations.history"
        f"?channel={urllib.parse.quote(channel)}&limit={limit}"
    )
    result = _http_get(url, headers=_bearer(token))
    if isinstance(result, dict) and not result.get("ok"):
        return f"Slack error: {result.get('error', result)}"
    messages = result.get("messages", [])
    lines = [f"[{m.get('ts', '')}] {m.get('text', '')}" for m in messages]
    return "\n".join(lines) if lines else "No messages found."


@tool(
    name="send_email",
    description=(
        "Send an email via SMTP. "
        "Requires SMTP_HOST, SMTP_USER, SMTP_PASSWORD env vars."
    ),
    risk=RiskTier.EXTERNAL_IO,
    tags=["communication", "email"],
)
async def send_email(to: str, subject: str, body: str) -> str:
    """Send a plain-text email using SMTP.

    to: Recipient email address.
    subject: Email subject line.
    body: Plain-text body content.
    """
    host = os.environ.get("SMTP_HOST", "")
    user = os.environ.get("SMTP_USER", "")
    password = os.environ.get("SMTP_PASSWORD", "")
    port = int(os.environ.get("SMTP_PORT", "587"))
    if not all([host, user, password]):
        return "Error: SMTP_HOST, SMTP_USER, and SMTP_PASSWORD must all be set"
    try:
        msg = MIMEMultipart()
        msg["From"] = user
        msg["To"] = to
        msg["Subject"] = subject
        msg.attach(MIMEText(body, "plain"))
        if port == 465:
            with SMTP_SSL(host, port, timeout=15) as smtp:
                smtp.login(user, password)
                smtp.sendmail(user, [to], msg.as_string())
        else:
            with SMTP(host, port, timeout=15) as smtp:
                smtp.ehlo()
                smtp.starttls()
                smtp.login(user, password)
                smtp.sendmail(user, [to], msg.as_string())
        return f"Email sent to {to}"
    except Exception as exc:
        return f"Error sending email: {exc}"


@tool(
    name="send_email_via_resend",
    description="Send an HTML email via the Resend API. Requires RESEND_API_KEY.",
    risk=RiskTier.EXTERNAL_IO,
    tags=["communication", "email"],
)
async def send_email_via_resend(to: str, subject: str, html: str) -> str:
    """Send a transactional HTML email using Resend.

    to: Recipient email address.
    subject: Email subject line.
    html: HTML body content.
    """
    api_key = os.environ.get("RESEND_API_KEY", "")
    if not api_key:
        return "Error: RESEND_API_KEY not set"
    from_addr = os.environ.get("RESEND_FROM", "noreply@example.com")
    result = _http_post(
        "https://api.resend.com/emails",
        {"from": from_addr, "to": [to], "subject": subject, "html": html},
        headers=_bearer(api_key),
    )
    if isinstance(result, dict) and "error" in result:
        return f"Resend error: {result['error']}"
    return f"Email queued via Resend (id={result.get('id', 'unknown')})"


@tool(
    name="teams_post_message",
    description="Post a message to a Microsoft Teams channel via an Incoming Webhook URL.",
    risk=RiskTier.EXTERNAL_IO,
    tags=["communication", "teams"],
)
async def teams_post_message(webhook_url: str, text: str) -> str:
    """Send a text message to a Teams channel using an Incoming Webhook.

    webhook_url: The full Incoming Webhook URL from Teams.
    text: Message body (plain text or Markdown).
    """
    result = _http_post(webhook_url, {"text": text})
    if isinstance(result, dict) and "error" in result:
        return f"Teams webhook error: {result['error']}"
    return "Message sent to Teams channel."


# ══════════════════════════════════════════════════════════════════════════════
# CATEGORY 2 — GITHUB
# ══════════════════════════════════════════════════════════════════════════════


def _gh_headers() -> dict[str, str]:
    token = os.environ.get("GITHUB_TOKEN", "")
    h = {"Accept": "application/vnd.github+json", "X-GitHub-Api-Version": "2022-11-28"}
    if token:
        h["Authorization"] = f"Bearer {token}"
    return h


@tool(
    name="github_get_repo",
    description="Get metadata for a GitHub repository. Requires GITHUB_TOKEN.",
    risk=RiskTier.READ_ONLY,
    tags=["github", "devops"],
)
async def github_get_repo(owner: str, repo: str) -> str:
    """Fetch repository metadata (stars, forks, description, default branch).

    owner: GitHub username or organization name.
    repo: Repository name.
    """
    result = _http_get(f"https://api.github.com/repos/{owner}/{repo}", headers=_gh_headers())
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    return json.dumps(
        {k: result.get(k) for k in ("full_name", "description", "stargazers_count", "forks_count", "default_branch", "open_issues_count")},
        indent=2,
    )


@tool(
    name="github_list_issues",
    description="List issues for a GitHub repository. Requires GITHUB_TOKEN.",
    risk=RiskTier.READ_ONLY,
    tags=["github", "devops"],
)
async def github_list_issues(owner: str, repo: str, state: str = "open") -> str:
    """Retrieve open or closed issues from a GitHub repository.

    owner: GitHub username or organization.
    repo: Repository name.
    state: 'open', 'closed', or 'all' (default 'open').
    """
    url = f"https://api.github.com/repos/{owner}/{repo}/issues?state={state}&per_page=20"
    result = _http_get(url, headers=_gh_headers())
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    if not isinstance(result, list):
        return f"Unexpected response: {result}"
    lines = [f"#{i.get('number')} [{i.get('state')}] {i.get('title')}" for i in result]
    return "\n".join(lines) if lines else "No issues found."


@tool(
    name="github_create_issue",
    description="Create a new issue in a GitHub repository. Requires GITHUB_TOKEN.",
    risk=RiskTier.INTERNAL,
    tags=["github", "devops"],
)
async def github_create_issue(owner: str, repo: str, title: str, body: str) -> str:
    """Open a new issue in a GitHub repository.

    owner: GitHub username or organization.
    repo: Repository name.
    title: Issue title.
    body: Issue body (Markdown supported).
    """
    token = os.environ.get("GITHUB_TOKEN", "")
    if not token:
        return "Error: GITHUB_TOKEN not set"
    result = _http_post(
        f"https://api.github.com/repos/{owner}/{repo}/issues",
        {"title": title, "body": body},
        headers=_gh_headers(),
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    return f"Issue created: #{result.get('number')} — {result.get('html_url')}"


@tool(
    name="github_get_file_content",
    description="Retrieve the content of a file from a GitHub repository. Requires GITHUB_TOKEN.",
    risk=RiskTier.READ_ONLY,
    tags=["github", "devops"],
)
async def github_get_file_content(owner: str, repo: str, path: str) -> str:
    """Fetch and decode a file from a GitHub repository.

    owner: GitHub username or organization.
    repo: Repository name.
    path: File path within the repository (e.g. 'src/main.py').
    """
    url = f"https://api.github.com/repos/{owner}/{repo}/contents/{path}"
    result = _http_get(url, headers=_gh_headers())
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    if isinstance(result, dict) and result.get("encoding") == "base64":
        try:
            return base64.b64decode(result["content"].replace("\n", "")).decode("utf-8")
        except Exception as exc:
            return f"Error decoding file: {exc}"
    return str(result)


@tool(
    name="github_search_code",
    description="Search code across GitHub repositories. Requires GITHUB_TOKEN.",
    risk=RiskTier.READ_ONLY,
    tags=["github", "devops"],
)
async def github_search_code(query: str) -> str:
    """Search for code matching a query across all public GitHub repositories.

    query: GitHub code search query string (e.g. 'meshflow tool:python').
    """
    url = f"https://api.github.com/search/code?q={urllib.parse.quote(query)}&per_page=10"
    result = _http_get(url, headers=_gh_headers())
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    items = result.get("items", [])
    lines = [f"{i.get('repository', {}).get('full_name')}: {i.get('path')}" for i in items]
    return f"Total: {result.get('total_count', 0)}\n" + "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# CATEGORY 3 — WEB SEARCH / SCRAPING
# ══════════════════════════════════════════════════════════════════════════════


@tool(
    name="web_search_serper",
    description="Search the web using the Serper.dev Google Search API. Requires SERPER_API_KEY.",
    risk=RiskTier.READ_ONLY,
    tags=["search", "web"],
)
async def web_search_serper(query: str, num: int = 5) -> str:
    """Perform a Google search via the Serper API and return organic results.

    query: Search query string.
    num: Number of results to return (default 5).
    """
    api_key = os.environ.get("SERPER_API_KEY", "")
    if not api_key:
        return "Error: SERPER_API_KEY not set"
    result = _http_post(
        "https://google.serper.dev/search",
        {"q": query, "num": num},
        headers={"X-API-KEY": api_key},
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    organic = result.get("organic", [])
    lines = [f"{r.get('title')}\n  {r.get('link')}\n  {r.get('snippet', '')}" for r in organic]
    return "\n\n".join(lines) if lines else "No results found."


@tool(
    name="web_search_tavily",
    description="Search the web using the Tavily AI Search API. Requires TAVILY_API_KEY.",
    risk=RiskTier.READ_ONLY,
    tags=["search", "web"],
)
async def web_search_tavily(query: str, max_results: int = 5) -> str:
    """Perform an AI-optimised web search using Tavily.

    query: Search query string.
    max_results: Maximum number of results (default 5).
    """
    api_key = os.environ.get("TAVILY_API_KEY", "")
    if not api_key:
        return "Error: TAVILY_API_KEY not set"
    result = _http_post(
        "https://api.tavily.com/search",
        {"api_key": api_key, "query": query, "max_results": max_results},
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    items = result.get("results", [])
    lines = [f"{r.get('title')}\n  {r.get('url')}\n  {r.get('content', '')[:200]}" for r in items]
    return "\n\n".join(lines) if lines else "No results found."


@tool(
    name="web_fetch_page",
    description="Fetch a web page and return its text content (HTML stripped).",
    risk=RiskTier.READ_ONLY,
    tags=["search", "web", "scraping"],
)
async def web_fetch_page(url: str, extract_text: bool = True) -> str:
    """Retrieve the contents of a URL, optionally stripping HTML tags.

    url: The URL to fetch.
    extract_text: If True (default), strip HTML tags and return clean text.
    """
    try:
        req = urllib.request.Request(
            url,
            headers={"User-Agent": "MeshFlow/1.0 (+https://meshflow.ai)"},
        )
        with urllib.request.urlopen(req, timeout=15) as resp:
            raw = resp.read().decode("utf-8", errors="replace")
    except Exception as exc:
        return f"Error fetching {url}: {exc}"

    if not extract_text:
        return raw[:8000]

    import re
    # Strip scripts and styles
    clean = re.sub(r"<(script|style)[^>]*>.*?</(script|style)>", "", raw, flags=re.DOTALL | re.IGNORECASE)
    # Strip HTML tags
    clean = re.sub(r"<[^>]+>", " ", clean)
    # Collapse whitespace
    clean = re.sub(r"\s+", " ", clean).strip()
    return clean[:8000]


@tool(
    name="web_get_sitemap",
    description="Fetch and parse the XML sitemap of a website.",
    risk=RiskTier.READ_ONLY,
    tags=["web", "scraping"],
)
async def web_get_sitemap(url: str) -> str:
    """Retrieve a website's sitemap.xml and return the list of URLs.

    url: Base URL of the site (e.g. 'https://example.com') or direct sitemap URL.
    """
    import re as _re

    sitemap_url = url if url.endswith(".xml") else url.rstrip("/") + "/sitemap.xml"
    try:
        req = urllib.request.Request(
            sitemap_url,
            headers={"User-Agent": "MeshFlow/1.0"},
        )
        with urllib.request.urlopen(req, timeout=15) as resp:
            raw = resp.read().decode("utf-8", errors="replace")
    except Exception as exc:
        return f"Error fetching sitemap: {exc}"

    urls = _re.findall(r"<loc>(.*?)</loc>", raw)
    return "\n".join(urls[:100]) if urls else "No URLs found in sitemap."


@tool(
    name="arxiv_search",
    description="Search academic papers on arXiv.org (no API key required).",
    risk=RiskTier.READ_ONLY,
    tags=["search", "research", "ai"],
)
async def arxiv_search(query: str, max_results: int = 5) -> str:
    """Search arXiv for academic papers matching a query.

    query: Search query string (e.g. 'transformer attention mechanism').
    max_results: Number of results to return (default 5, max 50).
    """
    import re as _re

    encoded = urllib.parse.quote(query)
    url = (
        f"https://export.arxiv.org/api/query?search_query=all:{encoded}"
        f"&start=0&max_results={max_results}"
    )
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "MeshFlow/1.0"})
        with urllib.request.urlopen(req, timeout=15) as resp:
            raw = resp.read().decode("utf-8", errors="replace")
    except Exception as exc:
        return f"Error querying arXiv: {exc}"

    entries = _re.findall(r"<entry>(.*?)</entry>", raw, _re.DOTALL)
    results = []
    for entry in entries[:max_results]:
        title_m = _re.search(r"<title>(.*?)</title>", entry, _re.DOTALL)
        id_m = _re.search(r"<id>(.*?)</id>", entry)
        summary_m = _re.search(r"<summary>(.*?)</summary>", entry, _re.DOTALL)
        title = (title_m.group(1) if title_m else "").strip().replace("\n", " ")
        arxiv_id = (id_m.group(1) if id_m else "").strip()
        summary = (summary_m.group(1) if summary_m else "").strip().replace("\n", " ")[:200]
        results.append(f"{title}\n  {arxiv_id}\n  {summary}")

    return "\n\n".join(results) if results else "No results found."


# ══════════════════════════════════════════════════════════════════════════════
# CATEGORY 4 — DATA / STORAGE
# ══════════════════════════════════════════════════════════════════════════════


@tool(
    name="postgres_query",
    description=(
        "Run a read or write SQL query against a PostgreSQL database. "
        "Requires DATABASE_URL env var. Requires psycopg2."
    ),
    risk=RiskTier.INTERNAL,
    tags=["database", "postgres"],
)
async def postgres_query(query: str, params: list[Any] | None = None) -> str:
    """Execute a SQL query against PostgreSQL and return results as JSON.

    query: SQL query string (SELECT, INSERT, UPDATE, DELETE, etc.).
    params: Optional list of query parameters for parameterised queries.
    """
    database_url = os.environ.get("DATABASE_URL", "")
    if not database_url:
        return "Error: DATABASE_URL not set"
    try:
        import psycopg2  # type: ignore[import]
        import psycopg2.extras  # type: ignore[import]
    except ImportError:
        return "Error: psycopg2 not installed — run: pip install psycopg2-binary"
    try:
        conn = psycopg2.connect(database_url)
        try:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                cur.execute(query, params or ())
                conn.commit()
                if cur.description:
                    rows = [dict(r) for r in cur.fetchmany(200)]
                    return json.dumps(rows, default=str)
                return f"Query OK — {cur.rowcount} row(s) affected."
        finally:
            conn.close()
    except Exception as exc:
        return f"Database error: {exc}"


@tool(
    name="redis_get",
    description="Get a value from Redis by key. Requires REDIS_URL env var.",
    risk=RiskTier.READ_ONLY,
    tags=["database", "redis", "cache"],
)
async def redis_get(key: str) -> str:
    """Retrieve a string value from Redis.

    key: The Redis key to look up.
    """
    redis_url = os.environ.get("REDIS_URL", "redis://localhost:6379")
    try:
        import redis  # type: ignore[import]
    except ImportError:
        return "Error: redis-py not installed — run: pip install redis"
    try:
        r = redis.from_url(redis_url, decode_responses=True)
        value = r.get(key)
        if value is None:
            return f"Key '{key}' not found."
        return str(value)
    except Exception as exc:
        return f"Redis error: {exc}"


@tool(
    name="redis_set",
    description="Set a key-value pair in Redis with optional TTL. Requires REDIS_URL env var.",
    risk=RiskTier.INTERNAL,
    tags=["database", "redis", "cache"],
)
async def redis_set(key: str, value: str, ttl: int = 3600) -> str:
    """Store a string value in Redis with a time-to-live.

    key: The Redis key.
    value: The string value to store.
    ttl: Time-to-live in seconds (default 3600 = 1 hour).
    """
    redis_url = os.environ.get("REDIS_URL", "redis://localhost:6379")
    try:
        import redis  # type: ignore[import]
    except ImportError:
        return "Error: redis-py not installed — run: pip install redis"
    try:
        r = redis.from_url(redis_url, decode_responses=True)
        r.set(key, value, ex=ttl)
        return f"Set '{key}' (TTL={ttl}s)"
    except Exception as exc:
        return f"Redis error: {exc}"


@tool(
    name="s3_get_object",
    description=(
        "Download an object from Amazon S3. "
        "Requires AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, AWS_DEFAULT_REGION."
    ),
    risk=RiskTier.READ_ONLY,
    tags=["storage", "aws", "s3"],
)
async def s3_get_object(bucket: str, key: str) -> str:
    """Retrieve the text content of an S3 object.

    bucket: S3 bucket name.
    key: Object key (path within the bucket).
    """
    try:
        import boto3  # type: ignore[import]
    except ImportError:
        return "Error: boto3 not installed — run: pip install boto3"
    try:
        s3 = boto3.client("s3")
        response = s3.get_object(Bucket=bucket, Key=key)
        body = response["Body"].read().decode("utf-8", errors="replace")
        return body[:8000]
    except Exception as exc:
        return f"S3 error: {exc}"


@tool(
    name="s3_list_objects",
    description=(
        "List objects in an Amazon S3 bucket. "
        "Requires AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, AWS_DEFAULT_REGION."
    ),
    risk=RiskTier.READ_ONLY,
    tags=["storage", "aws", "s3"],
)
async def s3_list_objects(bucket: str, prefix: str = "") -> str:
    """List object keys in an S3 bucket, optionally filtered by prefix.

    bucket: S3 bucket name.
    prefix: Key prefix filter (default '' = list all objects, up to 100).
    """
    try:
        import boto3  # type: ignore[import]
    except ImportError:
        return "Error: boto3 not installed — run: pip install boto3"
    try:
        s3 = boto3.client("s3")
        kwargs: dict[str, Any] = {"Bucket": bucket, "MaxKeys": 100}
        if prefix:
            kwargs["Prefix"] = prefix
        response = s3.list_objects_v2(**kwargs)
        objects = response.get("Contents", [])
        lines = [f"{o['Key']} ({o['Size']} bytes)" for o in objects]
        return "\n".join(lines) if lines else "No objects found."
    except Exception as exc:
        return f"S3 error: {exc}"


# ══════════════════════════════════════════════════════════════════════════════
# CATEGORY 5 — PRODUCTIVITY (Notion, Linear, Jira)
# ══════════════════════════════════════════════════════════════════════════════


@tool(
    name="notion_get_page",
    description="Retrieve a Notion page by ID. Requires NOTION_API_KEY.",
    risk=RiskTier.READ_ONLY,
    tags=["productivity", "notion"],
)
async def notion_get_page(page_id: str) -> str:
    """Fetch a Notion page's properties and plain-text content.

    page_id: The Notion page UUID (with or without dashes).
    """
    api_key = os.environ.get("NOTION_API_KEY", "")
    if not api_key:
        return "Error: NOTION_API_KEY not set"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Notion-Version": "2022-06-28",
    }
    result = _http_get(f"https://api.notion.com/v1/pages/{page_id}", headers=headers)
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    props = result.get("properties", {})
    title = ""
    for v in props.values():
        if isinstance(v, dict) and v.get("type") == "title":
            parts = v.get("title", [])
            title = "".join(p.get("plain_text", "") for p in parts)
            break
    return f"Page: {title}\nID: {result.get('id')}\nURL: {result.get('url')}"


@tool(
    name="notion_search",
    description="Search Notion workspace for pages and databases. Requires NOTION_API_KEY.",
    risk=RiskTier.READ_ONLY,
    tags=["productivity", "notion"],
)
async def notion_search(query: str) -> str:
    """Search across all Notion pages and databases accessible to the integration.

    query: Full-text search query.
    """
    api_key = os.environ.get("NOTION_API_KEY", "")
    if not api_key:
        return "Error: NOTION_API_KEY not set"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Notion-Version": "2022-06-28",
    }
    result = _http_post("https://api.notion.com/v1/search", {"query": query}, headers=headers)
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    items = result.get("results", [])
    lines = []
    for item in items[:10]:
        obj_type = item.get("object", "")
        url = item.get("url", "")
        props = item.get("properties", {})
        title = ""
        for v in props.values():
            if isinstance(v, dict) and v.get("type") == "title":
                parts = v.get("title", [])
                title = "".join(p.get("plain_text", "") for p in parts)
                break
        lines.append(f"[{obj_type}] {title or '(untitled)'} — {url}")
    return "\n".join(lines) if lines else "No results found."


@tool(
    name="linear_get_issues",
    description="List issues from a Linear team. Requires LINEAR_API_KEY.",
    risk=RiskTier.READ_ONLY,
    tags=["productivity", "linear"],
)
async def linear_get_issues(team_id: str, state: str = "Todo") -> str:
    """Fetch open issues for a Linear team filtered by state.

    team_id: Linear team ID (UUID format).
    state: Issue state name to filter by (default 'Todo').
    """
    api_key = os.environ.get("LINEAR_API_KEY", "")
    if not api_key:
        return "Error: LINEAR_API_KEY not set"
    gql = """
    query($teamId: String!, $filter: IssueFilter) {
      issues(filter: {team: {id: {eq: $teamId}}, state: {name: {eq: $filter}}}) {
        nodes { id title priority state { name } assignee { name } }
      }
    }
    """
    result = _http_post(
        "https://api.linear.app/graphql",
        {"query": gql, "variables": {"teamId": team_id, "filter": state}},
        headers=_bearer(api_key),
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    nodes = result.get("data", {}).get("issues", {}).get("nodes", [])
    lines = [
        f"[{n.get('priority')}] {n.get('title')} ({n.get('state', {}).get('name', '')})"
        for n in nodes
    ]
    return "\n".join(lines) if lines else "No issues found."


@tool(
    name="jira_get_issue",
    description="Get a Jira issue by key. Requires JIRA_URL, JIRA_USER, JIRA_TOKEN.",
    risk=RiskTier.READ_ONLY,
    tags=["productivity", "jira"],
)
async def jira_get_issue(issue_key: str) -> str:
    """Retrieve a Jira issue's summary, status, and description.

    issue_key: Jira issue key (e.g. 'PROJ-123').
    """
    jira_url = os.environ.get("JIRA_URL", "")
    user = os.environ.get("JIRA_USER", "")
    token = os.environ.get("JIRA_TOKEN", "")
    if not all([jira_url, user, token]):
        return "Error: JIRA_URL, JIRA_USER, and JIRA_TOKEN must all be set"
    credentials = base64.b64encode(f"{user}:{token}".encode()).decode()
    headers = {"Authorization": f"Basic {credentials}", "Accept": "application/json"}
    result = _http_get(f"{jira_url.rstrip('/')}/rest/api/3/issue/{issue_key}", headers=headers)
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    fields = result.get("fields", {})
    return (
        f"Key: {result.get('key')}\n"
        f"Summary: {fields.get('summary')}\n"
        f"Status: {fields.get('status', {}).get('name')}\n"
        f"Assignee: {(fields.get('assignee') or {}).get('displayName', 'Unassigned')}\n"
        f"Priority: {(fields.get('priority') or {}).get('name', 'None')}"
    )


@tool(
    name="jira_create_issue",
    description="Create a new Jira issue. Requires JIRA_URL, JIRA_USER, JIRA_TOKEN.",
    risk=RiskTier.INTERNAL,
    tags=["productivity", "jira"],
)
async def jira_create_issue(project: str, summary: str, description: str) -> str:
    """Create a new Jira issue (story/task) in a project.

    project: Jira project key (e.g. 'PROJ').
    summary: Issue summary/title.
    description: Detailed description (plain text).
    """
    jira_url = os.environ.get("JIRA_URL", "")
    user = os.environ.get("JIRA_USER", "")
    token = os.environ.get("JIRA_TOKEN", "")
    if not all([jira_url, user, token]):
        return "Error: JIRA_URL, JIRA_USER, and JIRA_TOKEN must all be set"
    credentials = base64.b64encode(f"{user}:{token}".encode()).decode()
    headers = {"Authorization": f"Basic {credentials}"}
    payload = {
        "fields": {
            "project": {"key": project},
            "summary": summary,
            "description": {
                "type": "doc",
                "version": 1,
                "content": [{"type": "paragraph", "content": [{"type": "text", "text": description}]}],
            },
            "issuetype": {"name": "Task"},
        }
    }
    result = _http_post(f"{jira_url.rstrip('/')}/rest/api/3/issue", payload, headers=headers)
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    return f"Issue created: {result.get('key')} — {jira_url.rstrip('/')}/browse/{result.get('key')}"


# ══════════════════════════════════════════════════════════════════════════════
# CATEGORY 6 — AI / LLM UTILITIES
# ══════════════════════════════════════════════════════════════════════════════


@tool(
    name="call_anthropic",
    description="Call the Anthropic Claude API with a prompt. Requires ANTHROPIC_API_KEY.",
    risk=RiskTier.READ_ONLY,
    tags=["ai", "llm", "anthropic"],
)
async def call_anthropic(
    prompt: str,
    model: str = "claude-haiku-4-5-20251001",
    max_tokens: int = 1000,
) -> str:
    """Send a prompt to Claude and return the text response.

    prompt: The user message to send.
    model: Anthropic model ID (default 'claude-haiku-4-5-20251001').
    max_tokens: Maximum tokens in the response (default 1000).
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return "Error: ANTHROPIC_API_KEY not set"
    result = _http_post(
        "https://api.anthropic.com/v1/messages",
        {
            "model": model,
            "max_tokens": max_tokens,
            "messages": [{"role": "user", "content": prompt}],
        },
        headers={"x-api-key": api_key, "anthropic-version": "2023-06-01"},
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    content = result.get("content", [])
    if content and isinstance(content, list):
        return content[0].get("text", "")
    return str(result)


@tool(
    name="call_openai",
    description="Call the OpenAI Chat Completions API with a prompt. Requires OPENAI_API_KEY.",
    risk=RiskTier.READ_ONLY,
    tags=["ai", "llm", "openai"],
)
async def call_openai(
    prompt: str,
    model: str = "gpt-4o-mini",
    max_tokens: int = 1000,
) -> str:
    """Send a prompt to OpenAI and return the text response.

    prompt: The user message to send.
    model: OpenAI model ID (default 'gpt-4o-mini').
    max_tokens: Maximum tokens in the response (default 1000).
    """
    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        return "Error: OPENAI_API_KEY not set"
    result = _http_post(
        "https://api.openai.com/v1/chat/completions",
        {
            "model": model,
            "max_tokens": max_tokens,
            "messages": [{"role": "user", "content": prompt}],
        },
        headers=_bearer(api_key),
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    choices = result.get("choices", [])
    if choices:
        return choices[0].get("message", {}).get("content", "")
    return str(result)


@tool(
    name="generate_embedding",
    description="Generate a text embedding vector using OpenAI. Requires OPENAI_API_KEY.",
    risk=RiskTier.READ_ONLY,
    tags=["ai", "embeddings", "openai"],
)
async def generate_embedding(text: str, model: str = "text-embedding-3-small") -> str:
    """Generate a numerical embedding vector for a text string.

    text: Input text to embed.
    model: OpenAI embedding model ID (default 'text-embedding-3-small').
    """
    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        return "Error: OPENAI_API_KEY not set"
    result = _http_post(
        "https://api.openai.com/v1/embeddings",
        {"input": text, "model": model},
        headers=_bearer(api_key),
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    data = result.get("data", [])
    if data:
        vec = data[0].get("embedding", [])
        return f"Embedding ({len(vec)} dims): [{', '.join(str(round(x, 6)) for x in vec[:5])}...]"
    return str(result)


@tool(
    name="moderate_content",
    description="Check text for harmful content using OpenAI Moderation API. Requires OPENAI_API_KEY.",
    risk=RiskTier.READ_ONLY,
    tags=["ai", "moderation", "safety"],
)
async def moderate_content(text: str) -> str:
    """Run text through OpenAI's content moderation endpoint.

    text: The text content to evaluate for harmful categories.
    """
    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        return "Error: OPENAI_API_KEY not set"
    result = _http_post(
        "https://api.openai.com/v1/moderations",
        {"input": text},
        headers=_bearer(api_key),
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    results = result.get("results", [])
    if results:
        r = results[0]
        flagged = r.get("flagged", False)
        categories = {k: v for k, v in r.get("categories", {}).items() if v}
        return f"Flagged: {flagged}\nTriggered categories: {list(categories.keys()) or 'none'}"
    return str(result)


@tool(
    name="translate_text",
    description="Translate text to a target language using OpenAI. Requires OPENAI_API_KEY.",
    risk=RiskTier.READ_ONLY,
    tags=["ai", "translation", "openai"],
)
async def translate_text(text: str, target_language: str) -> str:
    """Translate text into the specified target language using GPT.

    text: Text to translate.
    target_language: Target language name or BCP-47 code (e.g. 'Spanish', 'fr').
    """
    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        return "Error: OPENAI_API_KEY not set"
    result = _http_post(
        "https://api.openai.com/v1/chat/completions",
        {
            "model": "gpt-4o-mini",
            "max_tokens": 2000,
            "messages": [
                {
                    "role": "system",
                    "content": f"You are a professional translator. Translate the user's text to {target_language}. Output only the translated text.",
                },
                {"role": "user", "content": text},
            ],
        },
        headers=_bearer(api_key),
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    choices = result.get("choices", [])
    if choices:
        return choices[0].get("message", {}).get("content", "")
    return str(result)


# ══════════════════════════════════════════════════════════════════════════════
# CATEGORY 7 — CRM / SALES
# ══════════════════════════════════════════════════════════════════════════════


@tool(
    name="hubspot_get_contact",
    description="Look up a HubSpot contact by email address. Requires HUBSPOT_TOKEN.",
    risk=RiskTier.READ_ONLY,
    tags=["crm", "hubspot", "sales"],
)
async def hubspot_get_contact(email: str) -> str:
    """Fetch a HubSpot contact record by email address.

    email: The contact's email address.
    """
    token = os.environ.get("HUBSPOT_TOKEN", "")
    if not token:
        return "Error: HUBSPOT_TOKEN not set"
    url = (
        "https://api.hubapi.com/crm/v3/objects/contacts/search"
    )
    result = _http_post(
        url,
        {
            "filterGroups": [{"filters": [{"propertyName": "email", "operator": "EQ", "value": email}]}],
            "properties": ["firstname", "lastname", "email", "company", "phone"],
        },
        headers=_bearer(token),
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    items = result.get("results", [])
    if not items:
        return f"No contact found for {email}"
    c = items[0].get("properties", {})
    return f"Contact: {c.get('firstname')} {c.get('lastname')} | {c.get('email')} | {c.get('company')}"


@tool(
    name="hubspot_create_contact",
    description="Create a new contact in HubSpot. Requires HUBSPOT_TOKEN.",
    risk=RiskTier.INTERNAL,
    tags=["crm", "hubspot", "sales"],
)
async def hubspot_create_contact(email: str, firstname: str, lastname: str) -> str:
    """Create a new HubSpot contact record.

    email: Contact email address (must be unique).
    firstname: First name.
    lastname: Last name.
    """
    token = os.environ.get("HUBSPOT_TOKEN", "")
    if not token:
        return "Error: HUBSPOT_TOKEN not set"
    result = _http_post(
        "https://api.hubapi.com/crm/v3/objects/contacts",
        {"properties": {"email": email, "firstname": firstname, "lastname": lastname}},
        headers=_bearer(token),
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    return f"Contact created: {result.get('id')} — {email}"


@tool(
    name="salesforce_soql_query",
    description=(
        "Run a SOQL query against Salesforce. "
        "Requires SF_CLIENT_ID, SF_CLIENT_SECRET, SF_USERNAME, SF_PASSWORD."
    ),
    risk=RiskTier.READ_ONLY,
    tags=["crm", "salesforce", "sales"],
)
async def salesforce_soql_query(soql: str) -> str:
    """Execute a Salesforce Object Query Language (SOQL) statement.

    soql: SOQL query string (e.g. 'SELECT Id, Name FROM Account LIMIT 10').
    """
    client_id = os.environ.get("SF_CLIENT_ID", "")
    client_secret = os.environ.get("SF_CLIENT_SECRET", "")
    username = os.environ.get("SF_USERNAME", "")
    password = os.environ.get("SF_PASSWORD", "")
    if not all([client_id, client_secret, username, password]):
        return "Error: SF_CLIENT_ID, SF_CLIENT_SECRET, SF_USERNAME, SF_PASSWORD must all be set"
    # OAuth2 password flow
    token_data = urllib.parse.urlencode({
        "grant_type": "password",
        "client_id": client_id,
        "client_secret": client_secret,
        "username": username,
        "password": password,
    }).encode()
    try:
        req = urllib.request.Request(
            "https://login.salesforce.com/services/oauth2/token",
            data=token_data,
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=15) as resp:
            token_resp = json.loads(resp.read())
    except Exception as exc:
        return f"Salesforce auth error: {exc}"
    access_token = token_resp.get("access_token", "")
    instance_url = token_resp.get("instance_url", "")
    if not access_token:
        return f"Salesforce auth failed: {token_resp}"
    url = f"{instance_url}/services/data/v58.0/query?q={urllib.parse.quote(soql)}"
    result = _http_get(url, headers=_bearer(access_token))
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    records = result.get("records", [])
    return json.dumps(records[:50], default=str, indent=2)


@tool(
    name="pipedrive_get_deals",
    description="List deals from Pipedrive CRM. Requires PIPEDRIVE_TOKEN.",
    risk=RiskTier.READ_ONLY,
    tags=["crm", "pipedrive", "sales"],
)
async def pipedrive_get_deals(status: str = "open") -> str:
    """Retrieve deals from Pipedrive filtered by status.

    status: Deal status filter — 'open', 'won', 'lost', or 'all_not_deleted' (default 'open').
    """
    token = os.environ.get("PIPEDRIVE_TOKEN", "")
    if not token:
        return "Error: PIPEDRIVE_TOKEN not set"
    url = f"https://api.pipedrive.com/v1/deals?status={status}&api_token={token}&limit=20"
    result = _http_get(url)
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    data = result.get("data") or []
    lines = [f"[{d.get('status')}] {d.get('title')} — ${d.get('value', 0)}" for d in data]
    return "\n".join(lines) if lines else "No deals found."


@tool(
    name="clearbit_enrich_email",
    description="Enrich a person's profile from their email address using Clearbit. Requires CLEARBIT_KEY.",
    risk=RiskTier.READ_ONLY,
    tags=["crm", "enrichment", "sales"],
)
async def clearbit_enrich_email(email: str) -> str:
    """Look up person and company data for an email address via Clearbit Enrichment.

    email: The email address to enrich.
    """
    api_key = os.environ.get("CLEARBIT_KEY", "")
    if not api_key:
        return "Error: CLEARBIT_KEY not set"
    url = f"https://person.clearbit.com/v2/combined/find?email={urllib.parse.quote(email)}"
    result = _http_get(url, headers=_bearer(api_key))
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    person = result.get("person", {})
    company = result.get("company", {})
    name = f"{person.get('name', {}).get('givenName', '')} {person.get('name', {}).get('familyName', '')}".strip()
    return (
        f"Name: {name}\n"
        f"Title: {person.get('employment', {}).get('title', '')}\n"
        f"Company: {company.get('name', '')}\n"
        f"Domain: {company.get('domain', '')}\n"
        f"Industry: {company.get('category', {}).get('industry', '')}"
    )


# ══════════════════════════════════════════════════════════════════════════════
# CATEGORY 8 — CALENDAR / MEETINGS
# ══════════════════════════════════════════════════════════════════════════════


@tool(
    name="gcal_list_events",
    description="List upcoming Google Calendar events. Requires GOOGLE_CALENDAR_TOKEN.",
    risk=RiskTier.READ_ONLY,
    tags=["calendar", "google"],
)
async def gcal_list_events(calendar_id: str = "primary", max_results: int = 10) -> str:
    """Retrieve upcoming events from a Google Calendar.

    calendar_id: Calendar ID to query (default 'primary').
    max_results: Maximum number of events to return (default 10).
    """
    token = os.environ.get("GOOGLE_CALENDAR_TOKEN", "")
    if not token:
        return "Error: GOOGLE_CALENDAR_TOKEN not set"
    from datetime import datetime, timezone as _tz

    time_min = datetime.now(_tz.utc).isoformat()
    url = (
        f"https://www.googleapis.com/calendar/v3/calendars/{urllib.parse.quote(calendar_id)}/events"
        f"?maxResults={max_results}&singleEvents=true&orderBy=startTime&timeMin={urllib.parse.quote(time_min)}"
    )
    result = _http_get(url, headers=_bearer(token))
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    items = result.get("items", [])
    lines = [
        f"{e.get('start', {}).get('dateTime', e.get('start', {}).get('date'))} — {e.get('summary', '(no title)')}"
        for e in items
    ]
    return "\n".join(lines) if lines else "No upcoming events."


@tool(
    name="gcal_create_event",
    description="Create a Google Calendar event. Requires GOOGLE_CALENDAR_TOKEN.",
    risk=RiskTier.INTERNAL,
    tags=["calendar", "google"],
)
async def gcal_create_event(
    summary: str,
    start: str,
    end: str,
    attendees: list[str] | None = None,
) -> str:
    """Create an event in Google Calendar.

    summary: Event title.
    start: Start datetime in ISO 8601 format (e.g. '2025-01-15T10:00:00Z').
    end: End datetime in ISO 8601 format.
    attendees: Optional list of attendee email addresses.
    """
    token = os.environ.get("GOOGLE_CALENDAR_TOKEN", "")
    if not token:
        return "Error: GOOGLE_CALENDAR_TOKEN not set"
    payload: dict[str, Any] = {
        "summary": summary,
        "start": {"dateTime": start},
        "end": {"dateTime": end},
    }
    if attendees:
        payload["attendees"] = [{"email": a} for a in attendees]
    result = _http_post(
        "https://www.googleapis.com/calendar/v3/calendars/primary/events",
        payload,
        headers=_bearer(token),
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    return f"Event created: {result.get('summary')} — {result.get('htmlLink')}"


@tool(
    name="cal_com_get_bookings",
    description="List bookings from Cal.com. Requires CAL_COM_API_KEY.",
    risk=RiskTier.READ_ONLY,
    tags=["calendar", "cal.com"],
)
async def cal_com_get_bookings() -> str:
    """Retrieve upcoming bookings from your Cal.com account."""
    api_key = os.environ.get("CAL_COM_API_KEY", "")
    if not api_key:
        return "Error: CAL_COM_API_KEY not set"
    result = _http_get(
        f"https://api.cal.com/v1/bookings?apiKey={api_key}&status=upcoming"
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    bookings = result.get("bookings", [])
    lines = [
        f"{b.get('startTime')} — {b.get('title', '(no title)')} with {', '.join(a.get('email', '') for a in b.get('attendees', []))}"
        for b in bookings[:20]
    ]
    return "\n".join(lines) if lines else "No upcoming bookings."


@tool(
    name="zoom_list_meetings",
    description="List scheduled Zoom meetings. Requires ZOOM_TOKEN.",
    risk=RiskTier.READ_ONLY,
    tags=["calendar", "zoom"],
)
async def zoom_list_meetings() -> str:
    """Retrieve the list of upcoming Zoom meetings for the authenticated user."""
    token = os.environ.get("ZOOM_TOKEN", "")
    if not token:
        return "Error: ZOOM_TOKEN not set"
    result = _http_get(
        "https://api.zoom.us/v2/users/me/meetings?type=upcoming&page_size=20",
        headers=_bearer(token),
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    meetings = result.get("meetings", [])
    lines = [f"{m.get('start_time')} — {m.get('topic')} (id={m.get('id')})" for m in meetings]
    return "\n".join(lines) if lines else "No upcoming meetings."


@tool(
    name="zoom_create_meeting",
    description="Schedule a new Zoom meeting. Requires ZOOM_TOKEN.",
    risk=RiskTier.INTERNAL,
    tags=["calendar", "zoom"],
)
async def zoom_create_meeting(topic: str, duration_min: int = 60) -> str:
    """Create a new instant Zoom meeting.

    topic: Meeting topic/title.
    duration_min: Meeting duration in minutes (default 60).
    """
    token = os.environ.get("ZOOM_TOKEN", "")
    if not token:
        return "Error: ZOOM_TOKEN not set"
    result = _http_post(
        "https://api.zoom.us/v2/users/me/meetings",
        {"topic": topic, "type": 2, "duration": duration_min, "settings": {"join_before_host": True}},
        headers=_bearer(token),
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    return (
        f"Meeting created: {result.get('topic')}\n"
        f"Join URL: {result.get('join_url')}\n"
        f"Meeting ID: {result.get('id')}"
    )


# ══════════════════════════════════════════════════════════════════════════════
# CATEGORY 9 — FINANCE / DATA
# ══════════════════════════════════════════════════════════════════════════════


@tool(
    name="alpha_vantage_quote",
    description="Get a real-time stock quote from Alpha Vantage. Requires ALPHA_VANTAGE_KEY.",
    risk=RiskTier.READ_ONLY,
    tags=["finance", "stocks"],
)
async def alpha_vantage_quote(symbol: str) -> str:
    """Fetch the latest stock quote for a ticker symbol.

    symbol: Stock ticker symbol (e.g. 'AAPL', 'TSLA').
    """
    api_key = os.environ.get("ALPHA_VANTAGE_KEY", "")
    if not api_key:
        return "Error: ALPHA_VANTAGE_KEY not set"
    url = (
        f"https://www.alphavantage.co/query?function=GLOBAL_QUOTE"
        f"&symbol={urllib.parse.quote(symbol)}&apikey={api_key}"
    )
    result = _http_get(url)
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    quote = result.get("Global Quote", {})
    if not quote:
        return f"No data for symbol '{symbol}'"
    return (
        f"Symbol: {quote.get('01. symbol')}\n"
        f"Price: ${quote.get('05. price')}\n"
        f"Change: {quote.get('09. change')} ({quote.get('10. change percent')})\n"
        f"Volume: {quote.get('06. volume')}"
    )


@tool(
    name="coingecko_price",
    description="Get the current price of a cryptocurrency from CoinGecko (no API key required).",
    risk=RiskTier.READ_ONLY,
    tags=["finance", "crypto"],
)
async def coingecko_price(coin_id: str) -> str:
    """Fetch the current USD price and 24h change for a cryptocurrency.

    coin_id: CoinGecko coin ID (e.g. 'bitcoin', 'ethereum', 'solana').
    """
    url = (
        f"https://api.coingecko.com/api/v3/simple/price"
        f"?ids={urllib.parse.quote(coin_id)}&vs_currencies=usd&include_24hr_change=true"
    )
    result = _http_get(url)
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    data = result.get(coin_id, {})
    if not data:
        return f"No data found for '{coin_id}'"
    return (
        f"{coin_id}: ${data.get('usd', 'N/A')} USD "
        f"(24h change: {round(data.get('usd_24h_change', 0), 2)}%)"
    )


@tool(
    name="stripe_list_customers",
    description="List customers in Stripe. Requires STRIPE_SECRET_KEY.",
    risk=RiskTier.READ_ONLY,
    tags=["finance", "stripe", "payments"],
)
async def stripe_list_customers(limit: int = 10) -> str:
    """Retrieve recent Stripe customers.

    limit: Number of customers to return (default 10, max 100).
    """
    api_key = os.environ.get("STRIPE_SECRET_KEY", "")
    if not api_key:
        return "Error: STRIPE_SECRET_KEY not set"
    credentials = base64.b64encode(f"{api_key}:".encode()).decode()
    result = _http_get(
        f"https://api.stripe.com/v1/customers?limit={limit}",
        headers={"Authorization": f"Basic {credentials}"},
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    customers = result.get("data", [])
    lines = [f"{c.get('id')} — {c.get('email', '(no email)')} ({c.get('name', '')})" for c in customers]
    return "\n".join(lines) if lines else "No customers found."


@tool(
    name="stripe_create_payment_intent",
    description=(
        "Create a Stripe PaymentIntent to initiate a charge. "
        "Requires STRIPE_SECRET_KEY. IRREVERSIBLE financial operation."
    ),
    risk=RiskTier.IRREVERSIBLE,
    tags=["finance", "stripe", "payments"],
)
async def stripe_create_payment_intent(amount_cents: int, currency: str = "usd") -> str:
    """Create a Stripe PaymentIntent for a given amount.

    amount_cents: Amount to charge in the smallest currency unit (e.g. 1000 = $10.00 USD).
    currency: Three-letter ISO currency code (default 'usd').
    """
    api_key = os.environ.get("STRIPE_SECRET_KEY", "")
    if not api_key:
        return "Error: STRIPE_SECRET_KEY not set"
    credentials = base64.b64encode(f"{api_key}:".encode()).decode()
    data = urllib.parse.urlencode({"amount": str(amount_cents), "currency": currency}).encode()
    req = urllib.request.Request(
        "https://api.stripe.com/v1/payment_intents",
        data=data,
        headers={"Authorization": f"Basic {credentials}", "Content-Type": "application/x-www-form-urlencoded"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=15) as resp:
            result = json.loads(resp.read())
    except urllib.error.HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")
        return f"Stripe error: HTTP {exc.code}: {body[:400]}"
    except Exception as exc:
        return f"Stripe error: {exc}"
    return (
        f"PaymentIntent created: {result.get('id')}\n"
        f"Amount: {result.get('amount')} {result.get('currency')}\n"
        f"Status: {result.get('status')}\n"
        f"Client secret: {result.get('client_secret', '')[:20]}..."
    )


@tool(
    name="exchange_rate",
    description="Get the exchange rate between two currencies. Requires EXCHANGE_RATE_KEY.",
    risk=RiskTier.READ_ONLY,
    tags=["finance", "currency"],
)
async def exchange_rate(from_currency: str, to_currency: str) -> str:
    """Fetch the current exchange rate between two ISO 4217 currency codes.

    from_currency: Source currency code (e.g. 'USD').
    to_currency: Target currency code (e.g. 'EUR').
    """
    api_key = os.environ.get("EXCHANGE_RATE_KEY", "")
    if not api_key:
        return "Error: EXCHANGE_RATE_KEY not set"
    url = (
        f"https://v6.exchangerate-api.com/v6/{api_key}/pair"
        f"/{from_currency.upper()}/{to_currency.upper()}"
    )
    result = _http_get(url)
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    if result.get("result") == "error":
        return f"Exchange rate error: {result.get('error-type')}"
    rate = result.get("conversion_rate")
    return f"1 {from_currency.upper()} = {rate} {to_currency.upper()}"


# ══════════════════════════════════════════════════════════════════════════════
# CATEGORY 10 — DEVOPS / INFRASTRUCTURE
# ══════════════════════════════════════════════════════════════════════════════


@tool(
    name="github_actions_trigger",
    description="Trigger a GitHub Actions workflow dispatch event. Requires GITHUB_TOKEN.",
    risk=RiskTier.INTERNAL,
    tags=["devops", "github", "ci-cd"],
)
async def github_actions_trigger(
    owner: str, repo: str, workflow_id: str, ref: str = "main"
) -> str:
    """Trigger a GitHub Actions workflow via the workflow_dispatch event.

    owner: GitHub username or organization.
    repo: Repository name.
    workflow_id: Workflow file name or ID (e.g. 'deploy.yml').
    ref: Branch or tag to run the workflow on (default 'main').
    """
    token = os.environ.get("GITHUB_TOKEN", "")
    if not token:
        return "Error: GITHUB_TOKEN not set"
    result = _http_post(
        f"https://api.github.com/repos/{owner}/{repo}/actions/workflows/{workflow_id}/dispatches",
        {"ref": ref},
        headers=_gh_headers(),
        method="POST",
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    return f"Workflow '{workflow_id}' triggered on {owner}/{repo}@{ref}"


@tool(
    name="vercel_list_deployments",
    description="List recent Vercel deployments for a project. Requires VERCEL_TOKEN.",
    risk=RiskTier.READ_ONLY,
    tags=["devops", "vercel"],
)
async def vercel_list_deployments(project_id: str) -> str:
    """Retrieve the most recent deployments for a Vercel project.

    project_id: Vercel project ID or name.
    """
    token = os.environ.get("VERCEL_TOKEN", "")
    if not token:
        return "Error: VERCEL_TOKEN not set"
    url = f"https://api.vercel.com/v6/deployments?projectId={urllib.parse.quote(project_id)}&limit=10"
    result = _http_get(url, headers=_bearer(token))
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    deployments = result.get("deployments", [])
    lines = [
        f"{d.get('created', '')} [{d.get('state')}] {d.get('url')} ({d.get('meta', {}).get('githubCommitMessage', '')[:60]})"
        for d in deployments
    ]
    return "\n".join(lines) if lines else "No deployments found."


@tool(
    name="vercel_get_deployment",
    description="Get details of a specific Vercel deployment. Requires VERCEL_TOKEN.",
    risk=RiskTier.READ_ONLY,
    tags=["devops", "vercel"],
)
async def vercel_get_deployment(deployment_id: str) -> str:
    """Retrieve detailed information about a Vercel deployment.

    deployment_id: The Vercel deployment ID (e.g. 'dpl_xxxxxxxxxxxx').
    """
    token = os.environ.get("VERCEL_TOKEN", "")
    if not token:
        return "Error: VERCEL_TOKEN not set"
    result = _http_get(
        f"https://api.vercel.com/v13/deployments/{deployment_id}",
        headers=_bearer(token),
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    return (
        f"Deployment: {result.get('id')}\n"
        f"URL: {result.get('url')}\n"
        f"State: {result.get('readyState')}\n"
        f"Created: {result.get('createdAt')}\n"
        f"Creator: {result.get('creator', {}).get('email', '')}"
    )


@tool(
    name="datadog_post_event",
    description="Post a custom event to Datadog. Requires DD_API_KEY.",
    risk=RiskTier.INTERNAL,
    tags=["devops", "monitoring", "datadog"],
)
async def datadog_post_event(title: str, text: str, tags: list[str] | None = None) -> str:
    """Post a custom event to the Datadog events stream.

    title: Event title.
    text: Event body/description.
    tags: Optional list of tag strings (e.g. ['env:prod', 'service:api']).
    """
    api_key = os.environ.get("DD_API_KEY", "")
    if not api_key:
        return "Error: DD_API_KEY not set"
    payload: dict[str, Any] = {"title": title, "text": text}
    if tags:
        payload["tags"] = tags
    result = _http_post(
        "https://api.datadoghq.com/api/v1/events",
        payload,
        headers={"DD-API-KEY": api_key},
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    return f"Datadog event posted: {result.get('event', {}).get('id', 'ok')}"


@tool(
    name="pagerduty_create_incident",
    description=(
        "Create a PagerDuty incident. "
        "Requires PD_TOKEN. IRREVERSIBLE — triggers on-call alerts."
    ),
    risk=RiskTier.IRREVERSIBLE,
    tags=["devops", "pagerduty", "alerting"],
)
async def pagerduty_create_incident(
    title: str, service_id: str, urgency: str = "high"
) -> str:
    """Create a new PagerDuty incident and alert on-call responders.

    title: Incident title/description.
    service_id: PagerDuty service ID to associate with this incident.
    urgency: Incident urgency — 'high' or 'low' (default 'high').
    """
    token = os.environ.get("PD_TOKEN", "")
    email = os.environ.get("PD_FROM_EMAIL", "meshflow@example.com")
    if not token:
        return "Error: PD_TOKEN not set"
    payload = {
        "incident": {
            "type": "incident",
            "title": title,
            "service": {"id": service_id, "type": "service_reference"},
            "urgency": urgency,
        }
    }
    result = _http_post(
        "https://api.pagerduty.com/incidents",
        payload,
        headers={**_bearer(token), "From": email, "Accept": "application/vnd.pagerduty+json;version=2"},
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    incident = result.get("incident", {})
    return (
        f"Incident created: {incident.get('id')}\n"
        f"Title: {incident.get('title')}\n"
        f"Status: {incident.get('status')}\n"
        f"URL: {incident.get('html_url')}"
    )


# ══════════════════════════════════════════════════════════════════════════════
# CATEGORY 11 — FILE / DOCUMENT PROCESSING
# ══════════════════════════════════════════════════════════════════════════════


@tool(
    name="pdf_extract_text",
    description="Extract text from a PDF file (URL or local path). No API key required.",
    risk=RiskTier.READ_ONLY,
    tags=["documents", "pdf"],
)
async def pdf_extract_text(url_or_path: str) -> str:
    """Extract plain text from a PDF via URL or local file path.

    url_or_path: HTTP(S) URL or absolute local file path to the PDF.
    """
    try:
        import pypdf  # type: ignore[import]
    except ImportError:
        return "Error: pypdf not installed — run: pip install pypdf"
    import io

    try:
        if url_or_path.startswith("http://") or url_or_path.startswith("https://"):
            req = urllib.request.Request(url_or_path, headers={"User-Agent": "MeshFlow/1.0"})
            with urllib.request.urlopen(req, timeout=30) as resp:
                pdf_bytes = resp.read()
            reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
        else:
            reader = pypdf.PdfReader(url_or_path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text[:10000] if text else "No text extracted from PDF."
    except Exception as exc:
        return f"PDF extraction error: {exc}"


@tool(
    name="docx_extract_text",
    description="Extract text from a Word .docx file (URL or local path). No API key required.",
    risk=RiskTier.READ_ONLY,
    tags=["documents", "docx"],
)
async def docx_extract_text(url_or_path: str) -> str:
    """Extract plain text from a .docx file via URL or local file path.

    url_or_path: HTTP(S) URL or absolute local file path to the .docx file.
    """
    try:
        import docx  # type: ignore[import]
    except ImportError:
        return "Error: python-docx not installed — run: pip install python-docx"
    import io

    try:
        if url_or_path.startswith("http://") or url_or_path.startswith("https://"):
            req = urllib.request.Request(url_or_path, headers={"User-Agent": "MeshFlow/1.0"})
            with urllib.request.urlopen(req, timeout=30) as resp:
                docx_bytes = resp.read()
            doc = docx.Document(io.BytesIO(docx_bytes))
        else:
            doc = docx.Document(url_or_path)
        text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        return text[:10000] if text else "No text extracted from document."
    except Exception as exc:
        return f"DOCX extraction error: {exc}"


@tool(
    name="csv_parse",
    description="Parse CSV content from a URL or raw string and return JSON rows.",
    risk=RiskTier.READ_ONLY,
    tags=["documents", "csv", "data"],
)
async def csv_parse(url_or_content: str, delimiter: str = ",") -> str:
    """Parse CSV data from a URL or inline content string.

    url_or_content: HTTP(S) URL to a CSV file, or raw CSV text.
    delimiter: Field delimiter character (default ',').
    """
    import csv
    import io

    if url_or_content.startswith("http://") or url_or_content.startswith("https://"):
        try:
            req = urllib.request.Request(url_or_content, headers={"User-Agent": "MeshFlow/1.0"})
            with urllib.request.urlopen(req, timeout=15) as resp:
                content = resp.read().decode("utf-8", errors="replace")
        except Exception as exc:
            return f"Error fetching CSV: {exc}"
    else:
        content = url_or_content

    try:
        reader = csv.DictReader(io.StringIO(content), delimiter=delimiter)
        rows = list(reader)[:200]
        return json.dumps(rows, ensure_ascii=False)
    except Exception as exc:
        return f"CSV parse error: {exc}"


@tool(
    name="image_describe",
    description="Describe the content of an image at a URL using Claude Vision. Requires ANTHROPIC_API_KEY.",
    risk=RiskTier.READ_ONLY,
    tags=["documents", "vision", "ai"],
)
async def image_describe(url: str) -> str:
    """Generate a natural language description of an image using Claude's vision.

    url: Public URL of the image (JPEG, PNG, GIF, or WebP).
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return "Error: ANTHROPIC_API_KEY not set"
    # Determine media type from URL extension
    ext = url.rsplit(".", 1)[-1].lower()
    media_type_map = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png", "gif": "image/gif", "webp": "image/webp"}
    media_type = media_type_map.get(ext, "image/jpeg")
    result = _http_post(
        "https://api.anthropic.com/v1/messages",
        {
            "model": "claude-haiku-4-5-20251001",
            "max_tokens": 500,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "image", "source": {"type": "url", "url": url, "media_type": media_type}},
                        {"type": "text", "text": "Describe this image in detail."},
                    ],
                }
            ],
        },
        headers={"x-api-key": api_key, "anthropic-version": "2023-06-01"},
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    content = result.get("content", [])
    if content and isinstance(content, list):
        return content[0].get("text", "")
    return str(result)


@tool(
    name="ocr_image",
    description="Extract text from an image using Claude Vision OCR. Requires ANTHROPIC_API_KEY.",
    risk=RiskTier.READ_ONLY,
    tags=["documents", "ocr", "vision", "ai"],
)
async def ocr_image(url: str) -> str:
    """Perform OCR on an image at a URL using Claude's vision capability.

    url: Public URL of the image containing text to extract.
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return "Error: ANTHROPIC_API_KEY not set"
    ext = url.rsplit(".", 1)[-1].lower()
    media_type_map = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png", "gif": "image/gif", "webp": "image/webp"}
    media_type = media_type_map.get(ext, "image/jpeg")
    result = _http_post(
        "https://api.anthropic.com/v1/messages",
        {
            "model": "claude-haiku-4-5-20251001",
            "max_tokens": 1000,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "image", "source": {"type": "url", "url": url, "media_type": media_type}},
                        {"type": "text", "text": "Extract ALL text visible in this image. Output only the extracted text, preserving layout where possible."},
                    ],
                }
            ],
        },
        headers={"x-api-key": api_key, "anthropic-version": "2023-06-01"},
    )
    if isinstance(result, dict) and "error" in result:
        return result["error"]
    content = result.get("content", [])
    if content and isinstance(content, list):
        return content[0].get("text", "")
    return str(result)


# ══════════════════════════════════════════════════════════════════════════════
# REGISTRY EXPORT
# ══════════════════════════════════════════════════════════════════════════════

__all__ = [
    # Communication
    "slack_post_message",
    "slack_get_channel_history",
    "send_email",
    "send_email_via_resend",
    "teams_post_message",
    # GitHub
    "github_get_repo",
    "github_list_issues",
    "github_create_issue",
    "github_get_file_content",
    "github_search_code",
    # Web search / scraping
    "web_search_serper",
    "web_search_tavily",
    "web_fetch_page",
    "web_get_sitemap",
    "arxiv_search",
    # Data / Storage
    "postgres_query",
    "redis_get",
    "redis_set",
    "s3_get_object",
    "s3_list_objects",
    # Productivity
    "notion_get_page",
    "notion_search",
    "linear_get_issues",
    "jira_get_issue",
    "jira_create_issue",
    # AI / LLM utilities
    "call_anthropic",
    "call_openai",
    "generate_embedding",
    "moderate_content",
    "translate_text",
    # CRM / Sales
    "hubspot_get_contact",
    "hubspot_create_contact",
    "salesforce_soql_query",
    "pipedrive_get_deals",
    "clearbit_enrich_email",
    # Calendar / Meetings
    "gcal_list_events",
    "gcal_create_event",
    "cal_com_get_bookings",
    "zoom_list_meetings",
    "zoom_create_meeting",
    # Finance / Data
    "alpha_vantage_quote",
    "coingecko_price",
    "stripe_list_customers",
    "stripe_create_payment_intent",
    "exchange_rate",
    # DevOps / Infrastructure
    "github_actions_trigger",
    "vercel_list_deployments",
    "vercel_get_deployment",
    "datadog_post_event",
    "pagerduty_create_incident",
    # File / Document processing
    "pdf_extract_text",
    "docx_extract_text",
    "csv_parse",
    "image_describe",
    "ocr_image",
    # Registry
    "TOOL_REGISTRY",
]

TOOL_REGISTRY: dict[str, Any] = {
    # Communication
    "slack_post_message": slack_post_message,
    "slack_get_channel_history": slack_get_channel_history,
    "send_email": send_email,
    "send_email_via_resend": send_email_via_resend,
    "teams_post_message": teams_post_message,
    # GitHub
    "github_get_repo": github_get_repo,
    "github_list_issues": github_list_issues,
    "github_create_issue": github_create_issue,
    "github_get_file_content": github_get_file_content,
    "github_search_code": github_search_code,
    # Web search / scraping
    "web_search_serper": web_search_serper,
    "web_search_tavily": web_search_tavily,
    "web_fetch_page": web_fetch_page,
    "web_get_sitemap": web_get_sitemap,
    "arxiv_search": arxiv_search,
    # Data / Storage
    "postgres_query": postgres_query,
    "redis_get": redis_get,
    "redis_set": redis_set,
    "s3_get_object": s3_get_object,
    "s3_list_objects": s3_list_objects,
    # Productivity
    "notion_get_page": notion_get_page,
    "notion_search": notion_search,
    "linear_get_issues": linear_get_issues,
    "jira_get_issue": jira_get_issue,
    "jira_create_issue": jira_create_issue,
    # AI / LLM utilities
    "call_anthropic": call_anthropic,
    "call_openai": call_openai,
    "generate_embedding": generate_embedding,
    "moderate_content": moderate_content,
    "translate_text": translate_text,
    # CRM / Sales
    "hubspot_get_contact": hubspot_get_contact,
    "hubspot_create_contact": hubspot_create_contact,
    "salesforce_soql_query": salesforce_soql_query,
    "pipedrive_get_deals": pipedrive_get_deals,
    "clearbit_enrich_email": clearbit_enrich_email,
    # Calendar / Meetings
    "gcal_list_events": gcal_list_events,
    "gcal_create_event": gcal_create_event,
    "cal_com_get_bookings": cal_com_get_bookings,
    "zoom_list_meetings": zoom_list_meetings,
    "zoom_create_meeting": zoom_create_meeting,
    # Finance / Data
    "alpha_vantage_quote": alpha_vantage_quote,
    "coingecko_price": coingecko_price,
    "stripe_list_customers": stripe_list_customers,
    "stripe_create_payment_intent": stripe_create_payment_intent,
    "exchange_rate": exchange_rate,
    # DevOps / Infrastructure
    "github_actions_trigger": github_actions_trigger,
    "vercel_list_deployments": vercel_list_deployments,
    "vercel_get_deployment": vercel_get_deployment,
    "datadog_post_event": datadog_post_event,
    "pagerduty_create_incident": pagerduty_create_incident,
    # File / Document processing
    "pdf_extract_text": pdf_extract_text,
    "docx_extract_text": docx_extract_text,
    "csv_parse": csv_parse,
    "image_describe": image_describe,
    "ocr_image": ocr_image,
}
